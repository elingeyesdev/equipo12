from __future__ import annotations

import io
import re
from datetime import date
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


WORKSPACE = Path(r"C:\Univalle\equipo12")
SOURCE_DOCX = WORKSPACE / "manual_usuario_review.docx"
OUTPUT_DOCX = WORKSPACE / "Manual_Usuario_Carpooling_Univalle_Mejorado.docx"
MEDIA_DIR = WORKSPACE / "doc_review_media" / "manual_assets"

BLACK = RGBColor(0, 0, 0)
LIGHT_GRAY = "F2F2F2"
BORDER_GRAY = "BFBFBF"


FIGURES = {
    "login": ("image2.png", "Pantalla de inicio de sesion"),
    "register_basic": ("image3.png", "Registro inicial de usuario"),
    "passenger_home": ("image4.png", "Pantalla principal del pasajero"),
    "driver_found": ("image5.png", "Conductor encontrado para el trayecto"),
    "route_preview": ("image6.png", "Vista previa de ruta y decision de reserva"),
    "payment": ("image7.png", "Pantalla de pago de reserva"),
    "payment_check": ("image8.png", "Verificacion y estado del pago"),
    "boarding_code_menu": ("image9.png", "Menu principal y codigo de abordaje"),
    "register_vehicle": ("image10.png", "Registro de datos de conductor y vehiculo"),
    "create_trip": ("image11.png", "Creacion de viaje desde el mapa"),
    "fare": ("image12.png", "Ingreso del monto del viaje"),
    "trip_created": ("image13.png", "Viaje creado con acciones disponibles"),
    "driver_menu": ("image14.png", "Menu principal del conductor"),
    "passenger_request": ("image15.png", "Solicitud de pasajero pendiente"),
    "boarding_passenger": ("image16.png", "Seleccion de pasajero para abordaje"),
    "verification_code": ("image17.png", "Confirmacion mediante codigo de verificacion"),
    "route_progress": ("image18.jpeg", "Trayecto en curso"),
    "rating": ("image19.png", "Calificacion al finalizar el viaje"),
    "admin_login": ("image20.png", "Inicio de sesion del panel administrativo"),
    "admin_dashboard": ("image21.png", "Resumen y metricas del panel administrativo"),
    "admin_audit": ("image22.png", "Modulo de auditoria administrativa"),
    "admin_users": ("image23.png", "Administracion de usuarios"),
    "admin_safe_zones": ("image24.png", "Administracion de zonas seguras"),
}


def extract_media() -> None:
    MEDIA_DIR.mkdir(parents=True, exist_ok=True)
    with ZipFile(SOURCE_DOCX) as archive:
        for name in archive.namelist():
            if name.startswith("word/media/"):
                target = MEDIA_DIR / Path(name).name
                target.write_bytes(archive.read(name))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER_GRAY) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for side, value in values.items():
        element = tbl_cell_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def force_run_black(run) -> None:
    run.font.color.rgb = BLACK


def force_paragraph_black(paragraph) -> None:
    for run in paragraph.runs:
        force_run_black(run)


def force_document_black(doc: Document) -> None:
    for style in doc.styles:
        if getattr(style, "type", None) is not None:
            try:
                style.font.color.rgb = BLACK
            except Exception:
                pass
    for paragraph in doc.paragraphs:
        force_paragraph_black(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    force_paragraph_black(paragraph)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                force_paragraph_black(paragraph)


def apply_language_polish(doc: Document) -> None:
    replacements = {
        "Guia": "Guía",
        "guia": "guía",
        "actualizacion": "actualización",
        "Proposito": "Propósito",
        "proposito": "propósito",
        "Version": "Versión",
        "Descripcion": "Descripción",
        "descripcion": "descripción",
        "Indice": "Índice",
        "Introduccion": "Introducción",
        "introduccion": "introducción",
        "sesion": "sesión",
        "Sesion": "Sesión",
        "contrasena": "contraseña",
        "movil": "móvil",
        "aplicacion": "aplicación",
        "informacion": "información",
        "segun": "según",
        "operacion": "operación",
        "administracion": "administración",
        "administrativo": "administrativo",
        "funcion": "función",
        "Funcion": "Función",
        "vision": "visión",
        "autorizacion": "autorización",
        "conexion": "conexión",
        "ubicacion": "ubicación",
        "opcion": "opción",
        "presionar": "presionar",
        "Seleccionar": "Seleccionar",
        "seleccion": "selección",
        "Calificacion": "Calificación",
        "calificacion": "calificación",
        "codigo": "código",
        "Codigo": "Código",
        "verificacion": "verificación",
        "Verificacion": "Verificación",
        "creacion": "creación",
        "Creacion": "Creación",
        "decision": "decisión",
        "Decision": "Decisión",
        "Modulo": "Módulo",
        "modulo": "módulo",
        "Modulos": "Módulos",
        "modulos": "módulos",
        "metricas": "métricas",
        "graficos": "gráficos",
        "Auditoria": "Auditoría",
        "auditoria": "auditoría",
        "Pagina": "Página",
        "Que significa": "Qué significa",
        "Que hacer": "Qué hacer",
        "Que hago": "Qué hago",
        "esta": "está",
        "estan": "están",
        "esta dirigido": "está dirigido",
        "esta disponible": "está disponible",
        "esta intentando": "está intentando",
        "esta registrado": "está registrado",
        "esta opción": "esta opción",
        "validacion": "validación",
        "disponibilidad": "disponibilidad",
        "unicamente": "únicamente",
    }

    def polish_paragraph(paragraph) -> None:
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            original = text
            for source, target in replacements.items():
                pattern = r"(?<!\w)" + re.escape(source) + r"(?!\w)"
                text = re.sub(pattern, target, text)
            if text != original:
                run.text = text

    for paragraph in doc.paragraphs:
        polish_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    polish_paragraph(paragraph)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                polish_paragraph(paragraph)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, bold in [
        ("Title", 22, 0, 10, True),
        ("Subtitle", 12, 0, 10, False),
        ("Heading 1", 16, 18, 10, True),
        ("Heading 2", 13, 14, 7, True),
        ("Heading 3", 12, 10, 5, True),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = BLACK
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.font.color.rgb = BLACK
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(6)


def set_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for text in ("Pagina ",):
        run = paragraph.add_run(text)
        run.font.color.rgb = BLACK
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)
    run.font.color.rgb = BLACK


def add_text(doc: Document, text: str, style: str | None = None, align=None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.color.rgb = BLACK
    if align is not None:
        paragraph.alignment = align
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.color.rgb = BLACK


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.color.rgb = BLACK


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_cell_margins(table)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = BLACK
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.color.rgb = BLACK
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def image_size(path: Path, max_width_in: float, max_height_in: float) -> tuple[float, float]:
    with Image.open(path) as image:
        width_px, height_px = image.size
    ratio = width_px / height_px
    width = min(max_width_in, max_height_in * ratio)
    height = width / ratio
    if height > max_height_in:
        height = max_height_in
        width = height * ratio
    return width, height


def add_figure(doc: Document, key: str, number: int) -> int:
    filename, caption = FIGURES[key]
    path = MEDIA_DIR / filename
    if not path.exists():
        return number
    max_width = 3.25 if Image.open(path).height >= Image.open(path).width else 6.2
    max_height = 5.6 if Image.open(path).height >= Image.open(path).width else 3.2
    width, height = image_size(path, max_width, max_height)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    inline = run.add_picture(str(path), width=Inches(width), height=Inches(height))
    inline._inline.docPr.set("title", f"Figura {number}")
    inline._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(f"Figura {number}. {caption}. Fuente: Elaboracion propia.")
    cap_run.font.color.rgb = BLACK
    return number + 1


def add_section_title(doc: Document, title: str) -> None:
    add_text(doc, title, "Heading 1")


def add_subtitle(doc: Document, title: str) -> None:
    add_text(doc, title, "Heading 2")


def add_minor(doc: Document, title: str) -> None:
    add_text(doc, title, "Heading 3")


def create_manual() -> None:
    extract_media()
    doc = Document()
    set_page(doc.sections[0])
    configure_styles(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer)

    # Cover
    logo = MEDIA_DIR / "image1.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inline = p.add_run().add_picture(str(logo), width=Inches(1.35))
        inline._inline.docPr.set("title", "Logotipo institucional")
        inline._inline.docPr.set("descr", "Logotipo de la Universidad del Valle Bolivia")
    title = add_text(doc, "Manual de Usuario", "Title", WD_ALIGN_PARAGRAPH.CENTER)
    title.runs[0].font.size = Pt(24)
    add_text(doc, "Sistema de Carpooling Univalle", "Subtitle", WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Guia de uso para pasajeros, conductores y administradores", None, WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Equipo 12 - Proyecto de Sistemas", None, WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"Fecha de actualizacion: {date.today().strftime('%d/%m/%Y')}", None, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    add_section_title(doc, "Control del documento")
    add_table(
        doc,
        ["Dato", "Descripcion"],
        [
            ["Nombre del documento", "Manual de Usuario del Sistema de Carpooling Univalle"],
            ["Version", "1.0 revisada"],
            ["Audiencia", "Pasajeros, conductores, administradores y personal de soporte"],
            ["Proposito", "Explicar el uso correcto del sistema y orientar al usuario ante dudas o errores frecuentes"],
        ],
        [2.0, 4.5],
    )

    add_section_title(doc, "Indice general")
    add_numbered(
        doc,
        [
            "Introduccion",
            "Requisitos previos",
            "Roles del sistema",
            "Ingreso y registro de usuarios",
            "Uso como pasajero",
            "Uso como conductor",
            "Panel de control administrativo",
            "Manejo de errores y excepciones",
            "Buenas practicas de seguridad",
            "Preguntas frecuentes",
            "Glosario",
            "Indice de figuras",
        ],
    )
    doc.add_page_break()

    fig_no = 1

    add_section_title(doc, "1. Introduccion")
    add_text(
        doc,
        "Este manual explica, de manera ordenada y paso a paso, como utilizar el Sistema de Carpooling Univalle. "
        "El documento esta dirigido a los usuarios que ingresan desde la aplicacion movil y al personal que administra "
        "la informacion desde el panel web.",
    )
    add_text(
        doc,
        "El sistema permite que pasajeros soliciten viajes, conductores publiquen trayectos y administradores supervisen "
        "usuarios, reservas, pagos, zonas seguras, soporte y auditoria. Las capturas incluidas sirven como referencia visual; "
        "algunas pantallas pueden variar ligeramente segun el dispositivo, permisos o rol de usuario.",
    )

    add_subtitle(doc, "1.1 Alcance")
    add_bullets(
        doc,
        [
            "Crear una cuenta e iniciar sesion.",
            "Buscar, reservar, pagar y abordar un viaje como pasajero.",
            "Registrar vehiculo, crear viaje y gestionar pasajeros como conductor.",
            "Consultar informacion administrativa desde el panel web.",
            "Resolver errores frecuentes mediante mensajes claros y acciones recomendadas.",
        ],
    )

    add_section_title(doc, "2. Requisitos previos")
    add_bullets(
        doc,
        [
            "Contar con una cuenta institucional o los datos requeridos por Univalle para el registro.",
            "Tener conexion a internet activa.",
            "Permitir acceso a ubicacion en el celular para buscar rutas y mostrar el mapa correctamente.",
            "Usar un dispositivo Android compatible para la aplicacion movil.",
            "Para administradores, acceder desde un navegador web moderno y tener credenciales autorizadas.",
            "Verificar que el servidor y la base de datos del sistema se encuentren disponibles.",
        ],
    )

    add_section_title(doc, "3. Roles del sistema")
    add_table(
        doc,
        ["Rol", "Funciones principales"],
        [
            ["Pasajero", "Registrarse, buscar destino, solicitar reserva, pagar, abordar con codigo, revisar historial y calificar."],
            ["Conductor", "Registrar vehiculo, crear viajes, aceptar o rechazar reservas, validar codigos, iniciar recorrido y calificar pasajeros."],
            ["Administrador", "Gestionar usuarios, viajes, reservas, pagos, soporte, zonas seguras y reportes del sistema."],
            ["Superadministrador", "Administrar cuentas de administradores, roles, permisos y auditoria avanzada, cuando el sistema lo habilite."],
        ],
        [1.6, 4.9],
    )

    add_section_title(doc, "4. Ingreso y registro de usuarios")
    add_subtitle(doc, "4.1 Iniciar sesion")
    add_numbered(
        doc,
        [
            "Abrir la aplicacion movil o el panel web, segun corresponda.",
            "Ingresar el correo y la contrasena registrados.",
            "Presionar el boton de inicio de sesion.",
            "Si los datos son correctos, el sistema mostrara la pantalla principal correspondiente al rol del usuario.",
        ],
    )
    fig_no = add_figure(doc, "login", fig_no)

    add_subtitle(doc, "4.2 Registrar una cuenta")
    add_numbered(
        doc,
        [
            "Seleccionar la opcion de registro si aun no se cuenta con usuario.",
            "Completar los datos solicitados por el formulario.",
            "Elegir si el usuario utilizara el sistema como pasajero, conductor o ambos, cuando la pantalla lo permita.",
            "Revisar que la informacion sea correcta y presionar Registrar.",
        ],
    )
    fig_no = add_figure(doc, "register_basic", fig_no)

    add_section_title(doc, "5. Uso como pasajero")
    add_text(
        doc,
        "El pasajero utiliza la aplicacion para buscar un trayecto, revisar la informacion del conductor, reservar, realizar el pago "
        "y abordar el viaje mediante un codigo de verificacion.",
    )

    add_subtitle(doc, "5.1 Buscar destino")
    add_numbered(
        doc,
        [
            "Desde la pantalla principal, revisar la ubicacion actual mostrada en el mapa.",
            "Presionar la opcion de destino.",
            "Escribir o seleccionar el lugar al que se desea ir.",
            "Presionar Buscar conductor para ver opciones disponibles.",
        ],
    )
    fig_no = add_figure(doc, "passenger_home", fig_no)

    add_subtitle(doc, "5.2 Revisar conductor y ruta")
    add_numbered(
        doc,
        [
            "Leer la informacion del conductor y del trayecto sugerido.",
            "Revisar distancia, ruta y datos relevantes del viaje.",
            "Presionar Ver ruta para confirmar si el recorrido es conveniente.",
            "Aceptar la opcion si cumple con lo esperado o rechazarla para buscar otra alternativa.",
        ],
    )
    fig_no = add_figure(doc, "driver_found", fig_no)
    fig_no = add_figure(doc, "route_preview", fig_no)

    add_subtitle(doc, "5.3 Reservar y pagar el viaje")
    add_numbered(
        doc,
        [
            "Aceptar el viaje seleccionado.",
            "Esperar la confirmacion del conductor cuando la reserva lo requiera.",
            "Seleccionar el metodo de pago disponible.",
            "Presionar Continuar pago y verificar que el estado cambie correctamente.",
        ],
    )
    fig_no = add_figure(doc, "payment", fig_no)
    fig_no = add_figure(doc, "payment_check", fig_no)

    add_subtitle(doc, "5.4 Abordar el viaje")
    add_numbered(
        doc,
        [
            "Una vez confirmada la reserva, revisar el codigo de abordaje generado por la aplicacion.",
            "Entregar el codigo al conductor antes de iniciar el recorrido.",
            "Esperar a que el conductor confirme el abordaje.",
            "Mantener la aplicacion disponible durante el trayecto para revisar el estado del viaje.",
        ],
    )
    fig_no = add_figure(doc, "boarding_code_menu", fig_no)

    add_subtitle(doc, "5.5 Funciones adicionales del pasajero")
    add_table(
        doc,
        ["Funcion", "Uso recomendado"],
        [
            ["Historial de viajes", "Consultar viajes anteriores, revisar detalle del trayecto y verificar reservas completadas o canceladas."],
            ["Chat del viaje", "Comunicarse con el conductor cuando el sistema lo habilite para coordinar detalles del abordaje."],
            ["Favoritos", "Guardar ubicaciones frecuentes para buscar destinos con mayor rapidez."],
            ["Soporte", "Crear reportes cuando exista un problema con una reserva, un pago, un viaje o una cuenta."],
            ["Calificacion", "Valorar la experiencia al finalizar el recorrido para mejorar la confianza del sistema."],
        ],
        [1.8, 4.7],
    )

    add_section_title(doc, "6. Uso como conductor")
    add_text(
        doc,
        "El conductor registra sus datos y vehiculo, crea viajes, define el monto del trayecto, revisa solicitudes de pasajeros, valida "
        "codigos de abordaje e inicia el recorrido.",
    )

    add_subtitle(doc, "6.1 Registrar datos de conductor y vehiculo")
    add_numbered(
        doc,
        [
            "Ingresar a la opcion de registro.",
            "Completar los datos personales requeridos.",
            "Seleccionar la opcion de vehiculo si se desea operar como conductor.",
            "Completar placa, modelo, color y demas datos solicitados.",
            "Confirmar el registro.",
        ],
    )
    fig_no = add_figure(doc, "register_vehicle", fig_no)

    add_subtitle(doc, "6.2 Crear un viaje")
    add_numbered(
        doc,
        [
            "Desde la pantalla principal, seleccionar el destino en el mapa.",
            "Revisar el trazado de la ruta desde la ubicacion actual.",
            "Presionar Crear viaje.",
            "Ingresar el monto del viaje cuando el sistema lo solicite.",
            "Confirmar la creacion del viaje.",
        ],
    )
    fig_no = add_figure(doc, "create_trip", fig_no)
    fig_no = add_figure(doc, "fare", fig_no)
    fig_no = add_figure(doc, "trip_created", fig_no)

    add_subtitle(doc, "6.3 Gestionar solicitudes de pasajeros")
    add_numbered(
        doc,
        [
            "Abrir el menu principal del conductor.",
            "Ingresar a Solicitudes de pasajeros.",
            "Revisar la solicitud recibida.",
            "Aceptar o rechazar la solicitud segun disponibilidad y condiciones del viaje.",
            "Confirmar la decision para actualizar el estado de la reserva.",
        ],
    )
    fig_no = add_figure(doc, "driver_menu", fig_no)
    fig_no = add_figure(doc, "passenger_request", fig_no)

    add_subtitle(doc, "6.4 Confirmar abordaje e iniciar trayecto")
    add_numbered(
        doc,
        [
            "Seleccionar la opcion Abordar pasajero.",
            "Elegir al pasajero que realizo la reserva.",
            "Solicitar el codigo de abordaje mostrado en la aplicacion del pasajero.",
            "Ingresar el codigo y presionar Confirmar.",
            "Iniciar el recorrido cuando todos los pasajeros confirmados hayan abordado.",
        ],
    )
    fig_no = add_figure(doc, "boarding_passenger", fig_no)
    fig_no = add_figure(doc, "verification_code", fig_no)
    fig_no = add_figure(doc, "route_progress", fig_no)

    add_subtitle(doc, "6.5 Finalizar y calificar")
    add_numbered(
        doc,
        [
            "Finalizar el viaje al llegar al destino.",
            "Seleccionar la calificacion correspondiente.",
            "Escribir una observacion si se desea dejar mayor detalle.",
            "Enviar la calificacion para cerrar el proceso.",
        ],
    )
    fig_no = add_figure(doc, "rating", fig_no)

    add_subtitle(doc, "6.6 Funciones adicionales del conductor")
    add_table(
        doc,
        ["Funcion", "Descripcion"],
        [
            ["Pagos del conductor", "Consultar pagos relacionados con reservas y trayectos realizados."],
            ["Viajes programados", "Crear o revisar trayectos planificados cuando el sistema tenga horarios disponibles."],
            ["Historial", "Revisar viajes anteriores, pasajeros transportados y estados de reserva."],
            ["Soporte", "Reportar problemas de pasajeros, reservas, pagos o incidentes del recorrido."],
        ],
        [2.0, 4.5],
    )

    add_section_title(doc, "7. Panel de control administrativo")
    add_text(
        doc,
        "El panel administrativo permite supervisar la operacion del sistema. Solo los usuarios con permisos administrativos deben acceder "
        "a estas funciones.",
    )

    add_subtitle(doc, "7.1 Iniciar sesion en el panel")
    add_numbered(
        doc,
        [
            "Abrir el panel web del sistema.",
            "Ingresar el correo y contrasena de administrador.",
            "Presionar Entrar al panel.",
            "Verificar que se muestre el resumen administrativo.",
        ],
    )
    fig_no = add_figure(doc, "admin_login", fig_no)

    add_subtitle(doc, "7.2 Resumen y reportes")
    add_bullets(
        doc,
        [
            "Consultar cantidad de usuarios nuevos, viajes realizados y reservas generadas.",
            "Revisar graficos de actividad por dia, comparacion de pasajeros y conductores, y estado de reservas.",
            "Exportar reportes cuando se requiera respaldo en Excel, CSV o PDF.",
        ],
    )
    fig_no = add_figure(doc, "admin_dashboard", fig_no)

    add_subtitle(doc, "7.3 Modulos administrativos")
    add_table(
        doc,
        ["Modulo", "Acciones principales"],
        [
            ["Usuarios", "Buscar, revisar, editar, activar o desactivar cuentas segun permisos."],
            ["Viajes", "Consultar viajes registrados, rutas, conductores, pasajeros y estados."],
            ["Reservas", "Revisar reservas, estados, pasajeros asociados y fechas."],
            ["Pagos", "Supervisar pagos y transacciones relacionadas con las reservas."],
            ["Soporte", "Leer reportes de usuarios, responder mensajes y actualizar estados."],
            ["Zonas seguras", "Crear, editar o eliminar puntos seguros de recogida y destino."],
            ["Auditoria", "Revisar cambios administrativos, eventos sensibles y acciones relevantes."],
            ["Roles y administradores", "Crear cuentas administrativas y asignar permisos segun responsabilidad."],
            ["Personalizar", "Ajustar colores o tema visual del sistema si el rol lo permite."],
        ],
        [1.9, 4.6],
    )
    fig_no = add_figure(doc, "admin_audit", fig_no)
    fig_no = add_figure(doc, "admin_users", fig_no)
    fig_no = add_figure(doc, "admin_safe_zones", fig_no)

    add_section_title(doc, "8. Manejo de errores y excepciones")
    add_text(
        doc,
        "Cuando el sistema detecta un problema, muestra mensajes para orientar al usuario. La siguiente tabla resume los casos mas comunes "
        "y la accion recomendada.",
    )
    add_table(
        doc,
        ["Mensaje o situacion", "Que significa", "Que hacer"],
        [
            ["Usuario no encontrado", "La cuenta consultada no existe o fue eliminada.", "Verificar los datos ingresados o solicitar apoyo al administrador."],
            ["Vehiculo no encontrado", "El vehiculo no esta registrado en el sistema.", "Registrar nuevamente el vehiculo o revisar los datos guardados."],
            ["No se puede eliminar el vehiculo", "El vehiculo tiene viajes asociados.", "Eliminar, cerrar o reasignar los viajes relacionados antes de eliminarlo."],
            ["Viaje cancelado", "La operacion se esta intentando sobre un viaje que ya fue cancelado.", "Buscar otro viaje o revisar el estado actualizado."],
            ["El viaje no esta listo para iniciar", "Faltan datos, reservas o condiciones para iniciar.", "Completar la informacion requerida antes de presionar iniciar."],
            ["La tarifa no puede ser negativa", "El monto ingresado no es valido.", "Ingresar una tarifa mayor o igual a cero."],
            ["Coordenadas invalidas", "La ubicacion seleccionada no tiene formato valido.", "Seleccionar nuevamente el punto desde el mapa."],
            ["Categoria de soporte no valida", "El reporte fue creado con una categoria inexistente.", "Seleccionar una categoria disponible."],
            ["El mensaje no puede estar vacio", "Se intento enviar soporte o chat sin contenido.", "Escribir el mensaje antes de enviarlo."],
            ["Sesion expirada", "El acceso ya no es valido por seguridad.", "Cerrar la pantalla actual e iniciar sesion nuevamente."],
            ["No tienes permiso", "El usuario no cuenta con autorizacion para esa accion.", "Solicitar permisos o usar una cuenta autorizada."],
            ["No se pudo establecer conexion", "La app no pudo comunicarse con el servidor.", "Revisar internet, esperar unos minutos y volver a intentar."],
        ],
        [1.9, 2.2, 2.4],
    )

    add_section_title(doc, "9. Buenas practicas de seguridad")
    add_bullets(
        doc,
        [
            "No compartir contrasenas ni codigos de abordaje con personas ajenas al viaje.",
            "Verificar nombre del conductor, pasajero y ruta antes de confirmar una reserva.",
            "Usar zonas seguras para puntos de recogida y destino siempre que sea posible.",
            "Cerrar sesion en equipos compartidos.",
            "Reportar comportamientos sospechosos desde el modulo de soporte.",
            "Mantener actualizada la informacion del perfil y del vehiculo.",
        ],
    )

    add_section_title(doc, "10. Preguntas frecuentes")
    add_table(
        doc,
        ["Pregunta", "Respuesta"],
        [
            ["No puedo iniciar sesion. Que hago?", "Verifica correo, contrasena y conexion. Si el problema continua, solicita soporte o restablecimiento de acceso."],
            ["Como se confirma el abordaje?", "El pasajero muestra el codigo generado y el conductor lo ingresa para confirmar que abordo correctamente."],
            ["Puedo cancelar una reserva?", "Depende del estado de la reserva y de las reglas configuradas. Revisa el estado antes de cancelar."],
            ["Que hago si el pago no se confirma?", "Revisa la pantalla de estado del pago. Si continua pendiente, crea un reporte de soporte con el detalle de la reserva."],
            ["Como reporto un problema?", "Ingresa al modulo de soporte, selecciona la categoria correcta y describe el caso con la mayor claridad posible."],
        ],
        [2.4, 4.1],
    )

    add_section_title(doc, "11. Glosario")
    add_table(
        doc,
        ["Termino", "Definicion"],
        [
            ["Reserva", "Solicitud de espacio en un viaje creado por un conductor."],
            ["Abordaje", "Confirmacion de que el pasajero subio al vehiculo mediante codigo."],
            ["Zona segura", "Punto recomendado para recogida o destino dentro del sistema."],
            ["Tarifa", "Monto definido para el trayecto o reserva."],
            ["Auditoria", "Registro de acciones importantes realizadas en el sistema."],
            ["Soporte", "Canal para reportar problemas o solicitar ayuda."],
        ],
        [1.8, 4.7],
    )

    add_section_title(doc, "12. Indice de figuras")
    figure_rows = []
    for i, (_, caption) in enumerate(FIGURES.values(), 1):
        if i < fig_no:
            figure_rows.append([f"Figura {i}", caption])
    add_table(doc, ["Figura", "Descripcion"], figure_rows, [1.3, 5.2])

    add_section_title(doc, "13. Cierre")
    add_text(
        doc,
        "Este manual debe actualizarse cuando se agreguen nuevas pantallas, roles, permisos, reportes o cambios importantes en el flujo "
        "de viajes. Para mantenerlo util, cada nueva funcionalidad debe incluir una descripcion breve, pasos de uso, captura actualizada "
        "y posibles errores que el usuario podria encontrar.",
    )

    apply_language_polish(doc)
    force_document_black(doc)
    doc.core_properties.title = "Manual de Usuario - Sistema de Carpooling Univalle"
    doc.core_properties.subject = "Manual de usuario revisado"
    doc.core_properties.author = "Equipo 12"
    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    create_manual()
    print(OUTPUT_DOCX)
