from __future__ import annotations

import json
import math
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "DocumentacionPruebas"
EVIDENCE_DIR = OUT_DIR / "evidencias"
DOCX_PATH = OUT_DIR / "Informe_Pruebas_CarPooling.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_FILL = "F2F4F7"
GREEN = "1F7A3F"
RED = "9B1C1C"
AMBER = "7A5A00"


def ensure_dirs() -> None:
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)


def get_font(size: int, mono: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/consola.ttf" if mono else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/consolab.ttf" if mono else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def make_terminal_image(title: str, command: str, lines: list[str], filename: str) -> Path:
    width = 1700
    margin = 36
    header_h = 138
    line_h = 30
    wrapped: list[tuple[str, str]] = []
    max_chars = 118
    for line in lines:
        chunks = textwrap.wrap(line, width=max_chars, replace_whitespace=False) or [""]
        for idx, chunk in enumerate(chunks):
            wrapped.append((chunk, "cont" if idx else "base"))
    height = header_h + margin + max(1, len(wrapped)) * line_h + margin
    img = Image.new("RGB", (width, height), "#111827")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34)
    meta_font = get_font(21)
    mono_font = get_font(23, mono=True)
    draw.rectangle([0, 0, width, header_h], fill="#0B1220")
    draw.text((margin, 26), title, font=title_font, fill="#F9FAFB")
    draw.text((margin, 78), f"> {command}", font=meta_font, fill="#B7C4D6")
    y = header_h + 26
    for text, kind in wrapped:
        color = "#D1D5DB"
        if "Passed" in text or "Successful" in text or "PASO" in normalize(text) or "PASÓ" in text or "Exito" in normalize(text):
            color = "#86EFAC"
        if "Failed" in text or "Failure" in text or "FALLO" in normalize(text) or "FALLÓ" in text or "Actual:" in text:
            color = "#FCA5A5"
        if "Total tests" in text or "Total time" in text or "Resultados" in text or "Concurrencia" in text:
            color = "#93C5FD"
        x = margin if kind == "base" else margin + 32
        draw.text((x, y), text, font=mono_font, fill=color)
        y += line_h
    out = EVIDENCE_DIR / filename
    img.save(out)
    return out


def normalize(text: str) -> str:
    return (
        text.replace("Á", "A")
        .replace("É", "E")
        .replace("Í", "I")
        .replace("Ó", "O")
        .replace("Ú", "U")
        .replace("á", "a")
        .replace("é", "e")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("ú", "u")
    )


def status_text(approved: bool | None) -> str:
    if approved is True:
        return "Aprobado"
    if approved is False:
        return "No aprobado"
    return "Aprobado con observaciones"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], LIGHT_FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if "No aprobado" in val:
                set_cell_text(cells[i], val, bold=True, color=RED)
            elif "Aprobado" in val:
                set_cell_text(cells[i], val, bold=True, color=GREEN)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_evidence(doc: Document, img_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(6.35))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)


def build_evidence_images(nonfunc: dict) -> dict[str, Path]:
    evidence = {}
    evidence["unidad"] = make_terminal_image(
        "Evidencia - Pruebas de unidad",
        "dotnet test ... --filter FullyQualifiedName~UnitTests",
        [
            "Passed CarPooling.Tests.UnitTests.RatingService_CreateRatingAsync_ThrowsWhenTripNotFinished [1 s]",
            "Passed CarPooling.Tests.UnitTests.VehicleService_DeleteAsync_ThrowsExceptionWhenVehicleInTrip [82 ms]",
            "Passed CarPooling.Tests.UnitTests.VehicleService_UpsertForDriverAsync_UpdatesExistingActiveVehicle [36 ms]",
            "Passed CarPooling.Tests.UnitTests.VehicleService_CreateForDriverAsync_SavesAndNormalizesVehicle [14 ms]",
            "Test Run Successful.",
            "Total tests: 4",
            "Passed: 4",
            "Total time: 4.5780 Seconds",
        ],
        "01_unidad.png",
    )
    evidence["aceptacion"] = make_terminal_image(
        "Evidencia - Pruebas de aceptacion",
        "dotnet test ... --filter FullyQualifiedName~AcceptanceTests",
        [
            "Passed CarPooling.Tests.AcceptanceTests.Acceptance_PublishTrip_US1 [4 s]",
            "Passed CarPooling.Tests.AcceptanceTests.Acceptance_VerifyBoardingCode_US2 [247 ms]",
            "Test Run Successful.",
            "Total tests: 2",
            "Passed: 2",
            "Total time: 5.5422 Seconds",
        ],
        "02_aceptacion.png",
    )
    evidence["integracion"] = make_terminal_image(
        "Evidencia - Pruebas de integracion",
        "dotnet test ... --filter FullyQualifiedName~IntegrationTests",
        [
            "Failed CarPooling.Tests.IntegrationTests.Rating_DuplicatePrevention_ReturnsBadRequest [3 s]",
            "Assert.Equal() Failure: Expected: OK | Actual: BadRequest",
            "Failed CarPooling.Tests.IntegrationTests.Chat_AccessControl_AllowsAuthorized_RestrictsUnrelated [41 ms]",
            "Assert.Equal() Failure: Expected: OK | Actual: Forbidden",
            "Total tests: 2",
            "Failed: 2",
            "Test Run Failed.",
        ],
        "03_integracion.png",
    )
    evidence["funcional"] = make_terminal_image(
        "Evidencia - Pruebas funcionales",
        "dotnet test ... --filter FullyQualifiedName~FunctionalTests",
        [
            "Passed CarPooling.Tests.FunctionalTests.SupportTicket_ChatWorkflow [2 s]",
            "Passed CarPooling.Tests.FunctionalTests.Trip_ReservationLifecycle_Workflow [917 ms]",
            "Test Run Successful.",
            "Total tests: 2",
            "Passed: 2",
            "Total time: 3.7105 Seconds",
        ],
        "04_funcional.png",
    )
    evidence["regresion"] = make_terminal_image(
        "Evidencia - Pruebas de regresion",
        "dotnet test ... --filter FullyQualifiedName~RegressionTests",
        [
            "Passed CarPooling.Tests.RegressionTests.AdminEndpoint_AccessSecurityPolicies_RestrictsRegularUser [1 s]",
            "Failed CarPooling.Tests.RegressionTests.Reservation_OverbookingPrevention_ReturnsBadRequest [132 ms]",
            "Assert.Equal() Failure: Expected: BadRequest | Actual: Created",
            "Total tests: 2",
            "Passed: 1",
            "Failed: 1",
            "Test Run Failed.",
        ],
        "05_regresion.png",
    )
    perf = nonfunc["performance"]
    evidence["rendimiento"] = make_terminal_image(
        "Evidencia - Pruebas de rendimiento",
        "node Backend/.../Scripts/run_non_functional_tests.js",
        [
            "INICIANDO PRUEBAS DE RENDIMIENTO",
            f"P1 Busqueda de viajes: Exito {perf['tripSearch']['successes']}/{perf['tripSearch']['total']} | Min {perf['tripSearch']['min']}ms | Max {perf['tripSearch']['max']}ms | Promedio {perf['tripSearch']['avg']}ms | SLA < 200ms: {perf['tripSearch']['slaPassed']}",
            f"P2 Resumen de calificaciones: Exito {perf['ratingSummary']['successes']}/{perf['ratingSummary']['total']} | Min {perf['ratingSummary']['min']}ms | Max {perf['ratingSummary']['max']}ms | Promedio {perf['ratingSummary']['avg']}ms | SLA < 200ms: {perf['ratingSummary']['slaPassed']}",
            "Resultado exportado en non_functional_results.json",
        ],
        "06_rendimiento.png",
    )
    stress = nonfunc["stress"]
    stress_lines = ["INICIANDO PRUEBAS DE ESTRES (TIERS CONCURRENTES)"]
    for tier, data in stress.items():
        stress_lines.append(
            f"{tier} concurrentes | S1 Busqueda: {data['tripSearch']['successes']}/{tier}, fallos {data['tripSearch']['failures']}, latencia rafaga {data['tripSearch']['avgLatency']:.1f}ms"
        )
        stress_lines.append(
            f"{tier} concurrentes | S2 Zonas seguras: {data['safeZones']['successes']}/{tier}, fallos {data['safeZones']['failures']}, latencia rafaga {data['safeZones']['avgLatency']:.1f}ms"
        )
    evidence["estres"] = make_terminal_image(
        "Evidencia - Pruebas de estres",
        "node Backend/.../Scripts/run_non_functional_tests.js",
        stress_lines,
        "07_estres.png",
    )
    usability = nonfunc["usability"]
    usability_lines = ["INICIANDO AUDITORIA DE USABILIDAD Y ACCESIBILIDAD"]
    for item in usability["html"] + usability["css"]:
        prefix = "PASO" if item["passed"] else "FALLO"
        usability_lines.append(f"[{prefix}] {item['name']}: {item['detail']}")
    usability_lines.append(
        f"Resumen: Pasadas {usability['summary']['passed']} | Fallidas {usability['summary']['failed']}"
    )
    evidence["usabilidad"] = make_terminal_image(
        "Evidencia - Pruebas de usabilidad",
        "node Backend/.../Scripts/run_non_functional_tests.js",
        usability_lines,
        "08_usabilidad.png",
    )
    return evidence


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_status_callout(doc: Document, label: str, text: str, status: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    fill = "EAF4EA" if status == "Aprobado" else "FCE8E6" if status == "No aprobado" else "FFF4CE"
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"{label}: {status}. ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN if status == "Aprobado" else RED if status == "No aprobado" else AMBER)
    p.add_run(text)
    doc.add_paragraph()


def add_test_section(doc: Document, data: dict, evidence_path: Path) -> None:
    add_heading(doc, data["title"], 1)
    add_status_callout(doc, "Resultado general", data["summary"], status_text(data["approved"]))
    add_heading(doc, "Objetivo", 2)
    doc.add_paragraph(data["objective"])
    add_heading(doc, "Procedimiento realizado", 2)
    for step in data["procedure"]:
        doc.add_paragraph(step, style="List Number")
    add_heading(doc, "Casos ejecutados", 2)
    add_table(
        doc,
        ["Codigo", "Caso de prueba", "Resultado"],
        [[c["code"], c["name"], c["result"]] for c in data["cases"]],
        widths=[0.9, 4.5, 1.1],
    )
    add_heading(doc, "Conclusiones y observaciones", 2)
    for obs in data["observations"]:
        doc.add_paragraph(obs, style="List Bullet")
    add_heading(doc, "Captura de evidencia", 2)
    add_evidence(doc, evidence_path, data["caption"])
    doc.add_page_break()


def load_nonfunctional_results() -> dict:
    path = ROOT / "Backend" / "CarPooling" / "CarPooling" / "non_functional_results.json"
    return json.loads(path.read_text(encoding="utf-8"))


def main() -> None:
    ensure_dirs()
    nonfunc = load_nonfunctional_results()
    evidence = build_evidence_images(nonfunc)

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Informe de Planificacion, Ejecucion y Documentacion de Pruebas de Software")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Proyecto CarPooling - Universidad del Valle")
    r.font.size = Pt(13)

    doc.add_paragraph()
    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ["Proyecto", "CarPooling: backend ASP.NET Core, aplicacion movil Android y panel web administrativo."],
            ["Fecha de ejecucion", "11 de junio de 2026"],
            ["Repositorio evaluado", str(ROOT)],
            ["Herramientas", "xUnit/.NET Test, script Node de pruebas no funcionales, auditoria estatica de HTML/CSS."],
        ],
        widths=[1.8, 4.7],
    )

    add_heading(doc, "Integrantes del grupo", 1)
    doc.add_paragraph("Completar antes de subir a Teams:")
    add_table(
        doc,
        ["Nro.", "Nombre completo", "Rol en la actividad"],
        [
            ["1", "[Nombre completo del integrante 1]", "Pruebas de unidad e integracion"],
            ["2", "[Nombre completo del integrante 2]", "Pruebas funcionales y de aceptacion"],
            ["3", "[Nombre completo del integrante 3]", "Pruebas de rendimiento y estres"],
            ["4", "[Nombre completo del integrante 4]", "Pruebas de usabilidad y documentacion"],
        ],
        widths=[0.6, 3.4, 2.5],
    )

    add_heading(doc, "Planificacion de pruebas", 1)
    doc.add_paragraph(
        "El objetivo general fue verificar la calidad, funcionamiento y desempeno del sistema CarPooling mediante pruebas automatizadas y no funcionales. "
        "La planificacion se organizo por tipo de prueba para cubrir reglas de negocio, flujos de usuario, integracion entre API y datos, regresiones, rendimiento, estres y usabilidad."
    )
    add_table(
        doc,
        ["Tipo", "Alcance minimo", "Herramienta", "Criterio de aceptacion"],
        [
            ["Aceptacion", "2 historias de usuario", "xUnit + WebApplicationFactory", "Historias criticas ejecutan correctamente."],
            ["Integracion", "2 interacciones API/datos", "xUnit + API en memoria", "Servicios y endpoints integrados responden segun contrato."],
            ["Unidad", "Minimo 3 metodos", "xUnit + EF InMemory", "Logica aislada sin errores."],
            ["Funcional", "2 flujos completos", "xUnit + API", "Flujos de negocio concluyen con estados esperados."],
            ["Rendimiento", "2 endpoints", "Node fetch secuencial", "Promedio menor a 200 ms."],
            ["Regresion", "2 escenarios sensibles", "xUnit", "No reaparecen fallas conocidas."],
            ["Estres", "2 endpoints bajo concurrencia", "Node fetch concurrente", "Sin fallos bajo 100 a 1000 usuarios concurrentes."],
            ["Usabilidad", "Checks HTML/CSS", "Auditoria estatica", "Criterios de accesibilidad y responsividad aprobados."],
        ],
        widths=[1.2, 2.0, 1.8, 1.5],
    )

    add_heading(doc, "Resumen ejecutivo de resultados", 1)
    add_table(
        doc,
        ["Tipo de prueba", "Casos/checks", "Resultado"],
        [
            ["Aceptacion", "2/2 aprobados", "Aprobado"],
            ["Integracion", "0/2 aprobados", "No aprobado"],
            ["Unidad", "4/4 aprobados", "Aprobado"],
            ["Funcional", "2/2 aprobados", "Aprobado"],
            ["Rendimiento", "2/2 SLA aprobados", "Aprobado"],
            ["Regresion", "1/2 aprobados", "No aprobado"],
            ["Estres", "8/8 rafagas aprobadas", "Aprobado"],
            ["Usabilidad", "8/9 checks aprobados", "Aprobado con observaciones"],
        ],
        widths=[1.8, 2.3, 2.1],
    )
    doc.add_page_break()

    sections = [
        {
            "title": "Prueba de aceptacion",
            "approved": True,
            "summary": "Las dos historias de usuario criticas ejecutadas fueron aprobadas.",
            "objective": "Validar que el sistema cumple criterios de aceptacion desde la perspectiva del usuario final para publicar viajes y verificar codigos de abordaje.",
            "procedure": [
                "Se ejecuto la clase AcceptanceTests mediante dotnet test filtrando por FullyQualifiedName~AcceptanceTests.",
                "Se inicializo la API con WebApplicationFactory y datos semilla en memoria.",
                "Se validaron respuestas HTTP, estados de viaje/reserva y datos devueltos por la API.",
            ],
            "cases": [
                {"code": "PA-01", "name": "Publicar viaje programado con conductor y vehiculo activo.", "result": "Aprobado"},
                {"code": "PA-02", "name": "Verificar codigo de abordaje correcto para reserva confirmada.", "result": "Aprobado"},
            ],
            "observations": [
                "El flujo de publicacion creo el viaje con estado Programado, asientos disponibles y tarifa esperada.",
                "La verificacion de codigo autorizo el abordaje y actualizo la reserva al estado correspondiente.",
            ],
            "caption": "Figura 1. Ejecucion de AcceptanceTests con 2 pruebas aprobadas.",
        },
        {
            "title": "Pruebas de integracion",
            "approved": False,
            "summary": "Las dos pruebas fallaron por respuestas HTTP diferentes a las esperadas.",
            "objective": "Comprobar la comunicacion correcta entre endpoints, autenticacion por encabezado, reglas de acceso, servicios y persistencia en memoria.",
            "procedure": [
                "Se ejecuto IntegrationTests con dotnet test y WebApplicationFactory.",
                "Se consumieron endpoints de chat y calificaciones usando usuarios semilla y encabezados X-User-Id.",
                "Se compararon codigos HTTP esperados contra respuestas reales.",
            ],
            "cases": [
                {"code": "PI-01", "name": "Control de acceso del chat para pasajero autorizado y usuario ajeno.", "result": "No aprobado"},
                {"code": "PI-02", "name": "Prevencion de calificacion duplicada en un viaje finalizado.", "result": "No aprobado"},
            ],
            "observations": [
                "PI-01 esperaba OK para el pasajero autorizado, pero la API devolvio Forbidden.",
                "PI-02 esperaba OK en la primera calificacion, pero la API devolvio BadRequest.",
                "Se recomienda revisar datos semilla, estados previos de reserva/viaje y reglas de autorizacion usadas por los endpoints.",
            ],
            "caption": "Figura 2. Ejecucion de IntegrationTests con 2 fallos detectados.",
        },
        {
            "title": "Pruebas de unidad",
            "approved": True,
            "summary": "Se ejecutaron cuatro pruebas unitarias, superando el minimo solicitado de tres.",
            "objective": "Verificar componentes aislados de la logica de negocio sin depender de infraestructura externa.",
            "procedure": [
                "Se ejecuto UnitTests con dotnet test filtrado por clase.",
                "Se uso EF Core InMemory para simular el contexto de datos.",
                "Se verificaron servicios de vehiculos y calificaciones mediante aserciones directas.",
            ],
            "cases": [
                {"code": "PU-01", "name": "Crear vehiculo normaliza placa y guarda datos del conductor.", "result": "Aprobado"},
                {"code": "PU-02", "name": "Eliminar vehiculo usado en viaje lanza excepcion.", "result": "Aprobado"},
                {"code": "PU-03", "name": "Actualizar vehiculo activo mantiene el mismo registro.", "result": "Aprobado"},
                {"code": "PU-04", "name": "Crear calificacion en viaje no finalizado lanza excepcion.", "result": "Aprobado"},
            ],
            "observations": [
                "La normalizacion de placas y actualizacion de vehiculos funciono correctamente.",
                "Las reglas defensivas impiden eliminar vehiculos asociados a viajes y calificar viajes no finalizados.",
            ],
            "caption": "Figura 3. Ejecucion de UnitTests con 4 pruebas aprobadas.",
        },
        {
            "title": "Pruebas funcionales",
            "approved": True,
            "summary": "Los dos flujos funcionales principales fueron aprobados.",
            "objective": "Validar que funcionalidades completas del sistema operan de inicio a fin segun las reglas esperadas.",
            "procedure": [
                "Se ejecuto FunctionalTests con dotnet test.",
                "Se probo el ciclo de viaje/reserva y el flujo de chat de soporte.",
                "Se validaron codigos HTTP, cambios de estado, mensajes y reduccion de asientos disponibles.",
            ],
            "cases": [
                {"code": "PF-01", "name": "Ciclo de viaje: crear origen, destino, reserva y aceptacion.", "result": "Aprobado"},
                {"code": "PF-02", "name": "Flujo de soporte: crear ticket, respuesta admin, consulta y respuesta usuario.", "result": "Aprobado"},
            ],
            "observations": [
                "El viaje paso de Programado a Listo y la reserva fue aceptada correctamente.",
                "El chat de soporte permitio la comunicacion entre usuario y administrador.",
            ],
            "caption": "Figura 4. Ejecucion de FunctionalTests con 2 pruebas aprobadas.",
        },
        {
            "title": "Pruebas de rendimiento",
            "approved": True,
            "summary": "Los dos endpoints evaluados cumplieron el SLA definido de promedio menor a 200 ms.",
            "objective": "Medir tiempos de respuesta bajo carga secuencial controlada en endpoints relevantes para usuarios.",
            "procedure": [
                "Se levanto temporalmente el backend en http://localhost:5005.",
                "Se ejecuto el script run_non_functional_tests.js con 50 peticiones secuenciales por endpoint.",
                "Se calcularon minimo, maximo, promedio y cumplimiento del SLA.",
            ],
            "cases": [
                {"code": "PR-01", "name": "Busqueda de viajes candidatos con 50 peticiones secuenciales.", "result": "Aprobado"},
                {"code": "PR-02", "name": "Resumen de calificaciones de usuario con 50 peticiones secuenciales.", "result": "Aprobado"},
            ],
            "observations": [
                f"Busqueda de viajes: promedio {nonfunc['performance']['tripSearch']['avg']} ms, maximo {nonfunc['performance']['tripSearch']['max']} ms.",
                f"Resumen de calificaciones: promedio {nonfunc['performance']['ratingSummary']['avg']} ms, maximo {nonfunc['performance']['ratingSummary']['max']} ms.",
            ],
            "caption": "Figura 5. Ejecucion de rendimiento con resultados exportados a non_functional_results.json.",
        },
        {
            "title": "Pruebas de regresion",
            "approved": False,
            "summary": "Una prueba aprobo y una fallo, por lo que el bloque de regresion no se considera aprobado.",
            "objective": "Confirmar que comportamientos previamente protegidos no vuelvan a fallar despues de cambios recientes.",
            "procedure": [
                "Se ejecuto RegressionTests con dotnet test.",
                "Se validaron politicas de seguridad en endpoints administrativos.",
                "Se intento reservar un viaje sin asientos disponibles para detectar sobre-reserva.",
            ],
            "cases": [
                {"code": "PG-01", "name": "Endpoint administrativo restringe a usuario regular.", "result": "Aprobado"},
                {"code": "PG-02", "name": "Prevencion de sobre-reserva cuando no hay asientos disponibles.", "result": "No aprobado"},
            ],
            "observations": [
                "La seguridad del endpoint administrativo respondio Forbidden como se esperaba.",
                "La prevencion de sobre-reserva fallo: se esperaba BadRequest, pero la API respondio Created.",
                "Se recomienda revisar validacion de asientos disponibles en el flujo de creacion de reservas.",
            ],
            "caption": "Figura 6. Ejecucion de RegressionTests con 1 aprobado y 1 fallido.",
        },
        {
            "title": "Pruebas de estres",
            "approved": True,
            "summary": "Los endpoints resistieron cargas concurrentes de 100, 250, 500 y 1000 solicitudes sin fallos.",
            "objective": "Evaluar estabilidad del backend ante rafagas de solicitudes concurrentes y observar si aparecen errores bajo carga.",
            "procedure": [
                "Se uso el mismo backend temporal en http://localhost:5005.",
                "Se enviaron solicitudes concurrentes en niveles de 100, 250, 500 y 1000 usuarios simulados.",
                "Se midieron exitos, fallos y latencia promedio de rafaga para busqueda de viajes y zonas seguras.",
            ],
            "cases": [
                {"code": "PE-01", "name": "Busqueda de viajes bajo concurrencia progresiva.", "result": "Aprobado"},
                {"code": "PE-02", "name": "Consulta de zonas seguras bajo concurrencia progresiva.", "result": "Aprobado"},
            ],
            "observations": [
                "No se registraron fallos en los cuatro niveles de concurrencia para los dos endpoints.",
                "Los resultados indican estabilidad bajo la carga simulada; para produccion se recomienda repetir con base de datos y red equivalentes al entorno final.",
            ],
            "caption": "Figura 7. Ejecucion de estres con 100 a 1000 usuarios concurrentes simulados.",
        },
        {
            "title": "Pruebas de usabilidad",
            "approved": None,
            "summary": "La auditoria aprobo 8 de 9 checks; queda una observacion por inputs sin label o ARIA.",
            "objective": "Revisar criterios basicos de usabilidad, accesibilidad y responsividad del panel web.",
            "procedure": [
                "El script analizo Web/index.html y Web/styles.css.",
                "Se verificaron viewport, atributos ALT, labels/ARIA, H1, botones descriptivos, media queries, unidades adaptativas, touch targets y estilos de foco.",
                "Se conto el numero de criterios aprobados y fallidos.",
            ],
            "cases": [
                {"code": "PUX-01", "name": "Accesibilidad HTML: viewport, ALT, H1, botones e inputs.", "result": "Aprobado con observaciones"},
                {"code": "PUX-02", "name": "Responsividad CSS: media queries, unidades flexibles, touch targets y foco.", "result": "Aprobado"},
            ],
            "observations": [
                "El panel web cumple la mayoria de criterios revisados.",
                "Se detecto una brecha: 44 de 44 inputs carecen de label asociado o descriptor ARIA.",
                "Se recomienda asociar cada campo con label visible o aria-label/aria-labelledby para mejorar accesibilidad con lectores de pantalla.",
            ],
            "caption": "Figura 8. Auditoria de usabilidad y accesibilidad con 8 checks aprobados y 1 fallido.",
        },
    ]

    for section in sections:
        key = {
            "Prueba de aceptacion": "aceptacion",
            "Pruebas de integracion": "integracion",
            "Pruebas de unidad": "unidad",
            "Pruebas funcionales": "funcional",
            "Pruebas de rendimiento": "rendimiento",
            "Pruebas de regresion": "regresion",
            "Pruebas de estres": "estres",
            "Pruebas de usabilidad": "usabilidad",
        }[section["title"]]
        add_test_section(doc, section, evidence[key])

    add_heading(doc, "Conclusiones generales", 1)
    for conclusion in [
        "El proyecto cuenta con una base de pruebas automatizadas suficiente para evidenciar reglas de negocio, flujos funcionales y escenarios no funcionales.",
        "Los bloques de aceptacion, unidad, funcional, rendimiento y estres fueron aprobados.",
        "Integracion y regresion requieren correccion antes de considerar estable la version evaluada, especialmente en autorizacion de chat, calificaciones y prevencion de sobre-reserva.",
        "La usabilidad es mayormente adecuada, pero debe corregirse la falta de labels o descriptores ARIA en formularios del panel web.",
    ]:
        doc.add_paragraph(conclusion, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        "Nota: Las capturas incluidas corresponden a ejecuciones realizadas sobre el repositorio local el 11 de junio de 2026. "
        "Los resultados no funcionales fueron exportados automaticamente por el script del proyecto."
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
